"""
Lab part 1, program that creates an excel spreadsheet using data about countries from an API
Lab part 2, same thing but with Word
"""

import requests
from openpyxl import Workbook  # importing both excel library and api request library
import docx  # docx library for word file

url = 'https://country-list-1150.azurewebsites.net/api/country'  # assigning url

countries_response = requests.get(url).json()  # adding the api to the program. converting from json

country_workbook = Workbook()  # creating spreadsheet with openpyxl
worksheet = country_workbook.active  # assigning variable for the first page
worksheet.title = 'Countries and their Capitals'  # naming the page

row = 1  # assigning row that the data will start on

for dictionary in countries_response:  # looping over the list from the API, bunch of dictionaries with country data
    country_name = dictionary['name']  # assigning country name to variable from each dictionary
    capital_name = dictionary['capitalCity']  # assigning capital name to variable from each dictionary

    # print(f'{country_name}\'s capital is {capital_name}')

    worksheet.cell(row, 1, country_name)  # creating cell on the worksheet for each dictionary. uses variables
    worksheet.cell(row, 2, capital_name)  # that were just previously made
    row += 1  # add 1 to row for each time the loop runs, makes data continue to flow down spreadsheet.

country_workbook.save('countries_and_capitals.xlsx')  # saves the workbook

# ================================ Part 2 Word Document Creation ================================================

document = docx.Document()  # creating word document

document.add_paragraph('Countries and their Capital Cities', 'Title')  # adding title to document

for dictionary in countries_response:  # using loop in almost the exact same way as the spreadsheet, pulling from API.
    country_name = dictionary['name']
    capital_name = dictionary['capitalCity']  # assigning variables again otherwise it'd just use the last one from b4

    document.add_paragraph(f'The capital of {country_name} is {capital_name}')  # f-string using new variables and
    # .add_paragraph command from docx


document.save('countries_and_capitals.docx')  # saving document with the same name just different file type
