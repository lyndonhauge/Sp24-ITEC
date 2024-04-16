"""
program that creates a doc and uses a dictionary.
"""

import docx

document = docx.Document()

electric_cars = {
    'Chevrolet': 'Volt',
    'Nissan': 'Leaf',
    'Tesla': 'Model S',
    'Volkswagen': 'e-Golf'
}

document.add_paragraph('Electric Cars', 'Heading 1')

for make, model in electric_cars.items():  # loop over dictionary same as you normally would
    document.add_paragraph(make, 'Heading 3')
    document.add_paragraph(f'An electric car by {make} is {model}')

document.save('electric_cars.docx')
