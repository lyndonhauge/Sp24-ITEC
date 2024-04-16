"""
program that writes onto/creates a doc.
more types of styles you can use with docx inside the slides
"""

import docx

document = docx.Document()

document.add_paragraph('Hello Word!', 'Title')  # writing hello word inside doc and making it title

document.add_paragraph('By Lyndon', 'Heading 1')

document.add_paragraph('This is a word document created with Python and python-docx')

document.add_paragraph('Automate the boring stuff.', 'Quote')

document.add_paragraph('List of Favorite Colors', 'Heading 2')

favorite_colors = ['Blue', 'Purple', 'Green', 'Orange']
for color in favorite_colors:
    document.add_paragraph(color, 'List Bullet')


document.save('hello_word.docx')
