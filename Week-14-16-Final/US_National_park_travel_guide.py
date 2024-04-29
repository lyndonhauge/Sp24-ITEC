"""
Final project, program that creates a doc describing 5 random national parks with
data from an API
"""

import requests
import docx
import random
from pprint import pprint

park_list_url = 'https://national-parks-1150.azurewebsites.net/api/list'

park_list_response = requests.get(park_list_url).json()

random_parks_chosen = random.sample(park_list_response, 5)  # using .sample to pick a random 5 parks from list,
# storing them all inside a list

word_document = docx.Document()  # creating document
word_document.add_paragraph('National Park Travel Guide', 'Title')  # adding title

runs = 0

for park in random_parks_chosen:  # looping over the list of 5 random parks created to make more urls.

    park_code = park['park_code']  # these 3 lines are assigning park code to variable and putting it at the end of url
    # so that each park now has their own url which contains individual more specific information about them
    url_for_one_park_details = f'https://national-parks-1150.azurewebsites.net/api/{park_code}'
    # print(url_for_one_park_details)

    one_park_details_response = requests.get(url_for_one_park_details).json()  # response from each url
    # pprint(one_park_details_response)

    park_name = one_park_details_response['name']  # making variable for name to add to word doc
    word_document.add_paragraph(park_name, 'Heading 1')  # adding to doc

    description = one_park_details_response['description']  # making variable for description
    word_document.add_paragraph('Description', 'Heading 2')  # adding heading to doc
    word_document.add_paragraph(description)  # adding description data originally from api to doc

    weather = one_park_details_response['weather_overview']  # same idea again, just this time for weather
    word_document.add_paragraph('Weather', 'Heading 2')
    word_document.add_paragraph(weather)

    activities = one_park_details_response['activities']  # same again, but with loop to go over activities list.
    word_document.add_paragraph('Activities', 'Heading 2')
    for option in activities:
        word_document.add_paragraph(option, 'List Bullet')  # adding bullet points to doc for each activity

    images_data = one_park_details_response['nps_park_images']  # assigning variable and adding header
    word_document.add_paragraph('Images', 'Heading 2')
    num_of_images = len(images_data)  # counting the amount of items in nps_park_images list
    if num_of_images > 0:  # if there are images, proceed with this code
        for image in images_data:
            image_url = image['url']  # looping over list, assigning variables as needed
            image_caption = image['caption']
            image_credit = image['credit']
            image_response = requests.get(image_url)  # getting image response from url variable (for each image in loop
            with open('park_image.jpg', 'wb') as file:
                for chunk in image_response.iter_content():  # writing the file to save to directory each loop
                    file.write(chunk)  # it's fine if it overwrites because I can it add to docx once each run.
            word_document.add_picture('park_image.jpg', width=docx.shared.Inches(5.5))  # (adding saved image to docx)
            below_image = f'{image_caption} Credit: {image_credit}'  # credit and captions with variables made earlier
            word_document.add_paragraph(below_image, 'Quote')  # adding ^^^^^^^^^^^^^^^ to docx in quote format
    else:
        word_document.add_paragraph('No images available')  # if nothing is found in image list add this to docx

    operating_hours = one_park_details_response['park_operating_hours']  # assigning variable for key with data
    word_document.add_paragraph('Operating Hours', 'Heading 2')  # same type of format again
    word_document.add_paragraph(operating_hours)  # add the data below heading

    contact_info = one_park_details_response['contact_info']  # same type of process, just more assigning
    address = contact_info['address']  # assigning two more variables for nested dictionary data.
    website = contact_info['url']
    word_document.add_paragraph('Contact Information', 'Heading 2')  # printing out all that data onto the docx
    word_document.add_paragraph('Address', 'Heading 3')
    word_document.add_paragraph(address)
    word_document.add_paragraph('Website', 'Heading 3')
    word_document.add_paragraph(website)

    runs += 1
    print(f'National park {runs} done')

    # end of loop!


word_document.save('Park_Guide.docx')  # save once done
